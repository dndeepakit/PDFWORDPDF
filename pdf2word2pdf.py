import streamlit as st
from pdf2docx import Converter
import tempfile
import os
import traceback

st.set_page_config(page_title="PDF ↔ Word Converter", page_icon="📄", layout="centered")

st.title("📄 PDF ↔ Word Converter (Free Cloud Version)")
st.write("Convert between PDF and Word online — free and cloud-based.")

option = st.radio("Choose conversion type:", ["PDF → Word", "Word → PDF"])

uploaded_file = st.file_uploader("Upload your file", type=["pdf", "docx"])

if uploaded_file:
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        input_path = tmp.name

    try:
        if option == "PDF → Word" and suffix == ".pdf":
            output_path = input_path.replace(".pdf", ".docx")
            st.info("⏳ Converting your PDF to Word... please wait a moment.")
            try:
                cv = Converter(input_path)
                # Safe conversion – avoids get_area bug
                cv.convert(output_path, start=0, end=None)
                cv.close()
            except Exception as e:
                # Workaround for buggy pdf2docx Rect issues
                st.warning("⚠️ Falling back to safe mode...")
                try:
                    import fitz  # PyMuPDF
                    from docx import Document
                    doc = Document()
                    pdf = fitz.open(input_path)
                    for page in pdf:
                        text = page.get_text("text")
                        doc.add_paragraph(text)
                    doc.save(output_path)
                except Exception as e2:
                    st.error(f"Conversion failed even in fallback: {e2}")
                    st.stop()

            with open(output_path, "rb") as f:
                st.download_button("⬇️ Download Converted Word File", f, file_name="converted.docx")
            st.success("✅ Conversion completed successfully!")

        elif option == "Word → PDF" and suffix == ".docx":
            st.info("⏳ Converting your Word file to PDF...")
            from fpdf import FPDF
            from docx import Document
            doc = Document(input_path)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for para in doc.paragraphs:
                pdf.multi_cell(0, 10, para.text)
            output_path = input_path.replace(".docx", ".pdf")
            pdf.output(output_path)
            with open(output_path, "rb") as f:
                st.download_button("⬇️ Download Converted PDF", f, file_name="converted.pdf")
            st.success("✅ Conversion completed successfully!")
        else:
            st.warning("⚠️ Please upload the correct file type for your selected option.")

    except Exception as e:
        st.error(f"Conversion failed: {e}")
        st.text(traceback.format_exc())
